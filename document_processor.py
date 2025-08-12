#!/usr/bin/env python3
"""
Document Processor for Architecture Review Agent

Handles various document formats including Word documents, PDFs, and text files.
"""

import os
import re
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass
from enum import Enum
import mimetypes

# Import document processing libraries
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import docx2txt
    DOC_AVAILABLE = True
except ImportError:
    DOC_AVAILABLE = False

try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False


class DocumentFormat(Enum):
    """Supported document formats"""
    WORD_DOCX = "docx"
    WORD_DOC = "doc"
    PDF = "pdf"
    MARKDOWN = "md"
    TEXT = "txt"
    HTML = "html"
    RTF = "rtf"
    EXCEL = "xlsx"
    POWERPOINT = "pptx"
    UNKNOWN = "unknown"


@dataclass
class DocumentMetadata:
    """Document metadata extracted during processing"""
    format: DocumentFormat
    file_size: int
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    creation_date: Optional[str] = None
    last_modified: Optional[str] = None
    author: Optional[str] = None
    title: Optional[str] = None
    subject: Optional[str] = None
    sections: Optional[List[str]] = None
    tables_count: Optional[int] = None
    images_count: Optional[int] = None


class DocumentProcessor:
    """Processes various document formats and extracts text content"""
    
    def __init__(self):
        self.supported_formats = self._get_supported_formats()
        
    def _get_supported_formats(self) -> Dict[str, bool]:
        """Get supported formats based on available libraries"""
        return {
            'docx': DOCX_AVAILABLE,
            'doc': DOC_AVAILABLE,
            'pdf': PDF_AVAILABLE,
            'md': True,
            'txt': True,
            'html': True,
            'rtf': False,  # Requires additional library
            'xlsx': EXCEL_AVAILABLE,
            'pptx': False  # Requires additional library
        }
    
    def detect_format(self, file_path: str) -> DocumentFormat:
        """Detect document format based on file extension and content"""
        path = Path(file_path)
        extension = path.suffix.lower().lstrip('.')
        
        # First try by extension
        format_mapping = {
            'docx': DocumentFormat.WORD_DOCX,
            'doc': DocumentFormat.WORD_DOC,
            'pdf': DocumentFormat.PDF,
            'md': DocumentFormat.MARKDOWN,
            'markdown': DocumentFormat.MARKDOWN,
            'txt': DocumentFormat.TEXT,
            'text': DocumentFormat.TEXT,
            'html': DocumentFormat.HTML,
            'htm': DocumentFormat.HTML,
            'rtf': DocumentFormat.RTF,
            'xlsx': DocumentFormat.EXCEL,
            'xls': DocumentFormat.EXCEL,
            'pptx': DocumentFormat.POWERPOINT,
            'ppt': DocumentFormat.POWERPOINT
        }
        
        if extension in format_mapping:
            return format_mapping[extension]
        
        # Try MIME type detection
        mime_type, _ = mimetypes.guess_type(file_path)
        if mime_type:
            mime_mapping = {
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document': DocumentFormat.WORD_DOCX,
                'application/msword': DocumentFormat.WORD_DOC,
                'application/pdf': DocumentFormat.PDF,
                'text/markdown': DocumentFormat.MARKDOWN,
                'text/plain': DocumentFormat.TEXT,
                'text/html': DocumentFormat.HTML
            }
            if mime_type in mime_mapping:
                return mime_mapping[mime_type]
        
        return DocumentFormat.UNKNOWN
    
    def process_document(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Process a document and return its text content and metadata"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        doc_format = self.detect_format(file_path)
        
        if doc_format == DocumentFormat.WORD_DOCX:
            return self._process_docx(file_path)
        elif doc_format == DocumentFormat.WORD_DOC:
            return self._process_doc(file_path)
        elif doc_format == DocumentFormat.PDF:
            return self._process_pdf(file_path)
        elif doc_format == DocumentFormat.MARKDOWN:
            return self._process_markdown(file_path)
        elif doc_format == DocumentFormat.TEXT:
            return self._process_text(file_path)
        elif doc_format == DocumentFormat.HTML:
            return self._process_html(file_path)
        elif doc_format == DocumentFormat.EXCEL:
            return self._process_excel(file_path)
        else:
            raise ValueError(f"Unsupported document format: {doc_format.value}")
    
    def _process_docx(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Process Word DOCX document"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx library not available. Install with: pip install python-docx")
        
        doc = DocxDocument(file_path)
        
        # Extract text content
        text_parts = []
        sections = []
        tables_count = 0
        images_count = 0
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                # Check if it's a heading
                if paragraph.style.name.startswith('Heading'):
                    sections.append(text)
                text_parts.append(text)
        
        # Extract tables
        for table in doc.tables:
            tables_count += 1
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    table_text.append(" | ".join(row_text))
            
            if table_text:
                text_parts.append("\n" + "\n".join(table_text) + "\n")
        
        # Count images (inline shapes)
        try:
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    images_count += 1
        except:
            pass  # Image counting is optional
        
        full_text = "\n".join(text_parts)
        
        # Extract metadata
        props = doc.core_properties
        file_stats = os.stat(file_path)
        
        metadata = DocumentMetadata(
            format=DocumentFormat.WORD_DOCX,
            file_size=file_stats.st_size,
            word_count=len(full_text.split()),
            author=props.author,
            title=props.title,
            subject=props.subject,
            creation_date=str(props.created) if props.created else None,
            last_modified=str(props.modified) if props.modified else None,
            sections=sections,
            tables_count=tables_count,
            images_count=images_count
        )
        
        return full_text, metadata
    
    def _process_pdf(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Process PDF document"""
        if not PDF_AVAILABLE:
            raise ImportError("PDF processing libraries not available. Install with: pip install PyPDF2 pdfplumber")
        
        text_parts = []
        page_count = 0
        
        try:
            # Try pdfplumber first (better for complex layouts)
            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
        except:
            # Fallback to PyPDF2
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                page_count = len(reader.pages)
                
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
        
        full_text = "\n".join(text_parts)
        
        # Extract sections (simple heuristic based on common patterns)
        sections = self._extract_sections_from_text(full_text)
        
        file_stats = os.stat(file_path)
        
        metadata = DocumentMetadata(
            format=DocumentFormat.PDF,
            file_size=file_stats.st_size,
            page_count=page_count,
            word_count=len(full_text.split()),
            sections=sections,
            last_modified=str(file_stats.st_mtime)
        )
        
        return full_text, metadata
    
    def _process_markdown(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Process Markdown document"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Extract sections from markdown headers
        sections = []
        for match in re.finditer(r'^#+\s+(.+)$', content, re.MULTILINE):
            sections.append(match.group(1))
        
        file_stats = os.stat(file_path)
        
        metadata = DocumentMetadata(
            format=DocumentFormat.MARKDOWN,
            file_size=file_stats.st_size,
            word_count=len(content.split()),
            sections=sections,
            last_modified=str(file_stats.st_mtime)
        )
        
        return content, metadata
    
    def _process_text(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Process plain text document"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Extract potential sections
        sections = self._extract_sections_from_text(content)
        
        file_stats = os.stat(file_path)
        
        metadata = DocumentMetadata(
            format=DocumentFormat.TEXT,
            file_size=file_stats.st_size,
            word_count=len(content.split()),
            sections=sections,
            last_modified=str(file_stats.st_mtime)
        )
        
        return content, metadata
    
    def _process_html(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Process HTML document"""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("BeautifulSoup4 not available. Install with: pip install beautifulsoup4")
        
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        soup = BeautifulSoup(content, 'html.parser')
        
        # Extract text content
        text = soup.get_text()
        
        # Extract sections from headings
        sections = []
        for heading in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
            if heading.text.strip():
                sections.append(heading.text.strip())
        
        # Count tables and images
        tables_count = len(soup.find_all('table'))
        images_count = len(soup.find_all('img'))
        
        file_stats = os.stat(file_path)
        
        metadata = DocumentMetadata(
            format=DocumentFormat.HTML,
            file_size=file_stats.st_size,
            word_count=len(text.split()),
            sections=sections,
            tables_count=tables_count,
            images_count=images_count,
            title=soup.title.string if soup.title else None,
            last_modified=str(file_stats.st_mtime)
        )
        
        return text, metadata
    
    def _process_excel(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Process Excel document (extract text from worksheets)"""
        if not EXCEL_AVAILABLE:
            raise ImportError("openpyxl not available. Install with: pip install openpyxl")
        
        workbook = openpyxl.load_workbook(file_path, data_only=True)
        
        text_parts = []
        sections = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sections.append(f"Sheet: {sheet_name}")
            
            sheet_text = []
            for row in sheet.iter_rows():
                row_text = []
                for cell in row:
                    if cell.value is not None:
                        row_text.append(str(cell.value))
                if row_text:
                    sheet_text.append(" | ".join(row_text))
            
            if sheet_text:
                text_parts.append(f"\n=== {sheet_name} ===\n" + "\n".join(sheet_text))
        
        full_text = "\n".join(text_parts)
        
        file_stats = os.stat(file_path)
        
        metadata = DocumentMetadata(
            format=DocumentFormat.EXCEL,
            file_size=file_stats.st_size,
            word_count=len(full_text.split()),
            sections=sections,
            tables_count=len(workbook.sheetnames),
            last_modified=str(file_stats.st_mtime)
        )
        
        return full_text, metadata
    
    def _process_doc(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Process older Word DOC document"""
        if not DOC_AVAILABLE:
            raise ImportError("docx2txt library not available. Install with: pip install python-docx2txt")
        
        # Extract text content using docx2txt
        content = docx2txt.process(file_path)
        
        # Extract potential sections from text
        sections = self._extract_sections_from_text(content)
        
        # Get file statistics
        file_stats = os.stat(file_path)
        
        metadata = DocumentMetadata(
            format=DocumentFormat.WORD_DOC,
            file_size=file_stats.st_size,
            word_count=len(content.split()),
            sections=sections,
            last_modified=str(file_stats.st_mtime)
        )
        
        return content, metadata
    
    def _extract_sections_from_text(self, text: str) -> List[str]:
        """Extract potential sections from plain text using common patterns"""
        sections = []
        
        # Common section patterns
        patterns = [
            r'^([A-Z][A-Za-z\s]+):?\s*$',  # Title case lines
            r'^\d+\.\s+([A-Z][A-Za-z\s]+)',  # Numbered sections
            r'^([A-Z\s]{3,})\s*$',  # ALL CAPS headers
            r'^-+\s*([A-Za-z\s]+)\s*-+$',  # Dashed headers
            r'^=+\s*([A-Za-z\s]+)\s*=+$',  # Equals headers
        ]
        
        for pattern in patterns:
            matches = re.finditer(pattern, text, re.MULTILINE)
            for match in matches:
                section = match.group(1).strip()
                if len(section) > 3 and section not in sections:
                    sections.append(section)
        
        return sections[:20]  # Limit to first 20 sections
    
    def get_installation_instructions(self) -> Dict[str, str]:
        """Get installation instructions for missing libraries"""
        instructions = {}
        
        if not DOCX_AVAILABLE:
            instructions['Word Documents (.docx)'] = "pip install python-docx"
        
        if not DOC_AVAILABLE:
            instructions['Word Documents (.doc)'] = "pip install python-docx2txt"
        
        if not PDF_AVAILABLE:
            instructions['PDF Documents'] = "pip install PyPDF2 pdfplumber"
        
        if not EXCEL_AVAILABLE:
            instructions['Excel Documents'] = "pip install openpyxl"
        
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            instructions['HTML Documents'] = "pip install beautifulsoup4"
        
        return instructions
    
    def check_dependencies(self) -> Dict[str, bool]:
        """Check which document processing dependencies are available"""
        dependencies = {
            'python-docx (Word .docx)': DOCX_AVAILABLE,
            'docx2txt (Word .doc)': DOC_AVAILABLE,
            'PyPDF2/pdfplumber (PDF)': PDF_AVAILABLE,
            'openpyxl (Excel)': EXCEL_AVAILABLE,
            'pandas (Data Analysis)': PANDAS_AVAILABLE
        }
        
        try:
            from bs4 import BeautifulSoup
            dependencies['beautifulsoup4 (HTML)'] = True
        except ImportError:
            dependencies['beautifulsoup4 (HTML)'] = False
        
        return dependencies


def main():
    """Demo the document processor"""
    processor = DocumentProcessor()
    
    print("📄 Document Processor Demo")
    print("=" * 50)
    
    # Check dependencies
    print("🔍 Checking Dependencies:")
    deps = processor.check_dependencies()
    for name, available in deps.items():
        status = "✅ Available" if available else "❌ Missing"
        print(f"  {name}: {status}")
    
    # Installation instructions
    missing = processor.get_installation_instructions()
    if missing:
        print(f"\n📦 To install missing dependencies:")
        for doc_type, command in missing.items():
            print(f"  {doc_type}: {command}")
    
    print(f"\n📊 Supported Formats:")
    for fmt, supported in processor.supported_formats.items():
        status = "✅" if supported else "❌"
        print(f"  .{fmt}: {status}")


if __name__ == "__main__":
    main()
