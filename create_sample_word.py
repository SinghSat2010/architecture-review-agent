#!/usr/bin/env python3
"""
Create Sample Word Document

This script creates a sample Word document (.docx) for testing
the document processing capabilities.
"""

try:
    from docx import Document
    from docx.shared import Inches
    
    # Create a new document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Sample Architecture Document', 0)
    
    # Add executive summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph('This document outlines the architecture for a high-availability, scalable system designed to meet enterprise requirements.')
    
    # Add business requirements
    doc.add_heading('Business Requirements', level=1)
    business_reqs = doc.add_paragraph()
    business_reqs.add_run('The system must support:').bold = True
    business_reqs.add_run('\n• High availability (99.9% uptime)\n• Scalable architecture\n• Security compliance\n• Cost optimization')
    
    # Add technical requirements
    doc.add_heading('Technical Requirements', level=1)
    tech_reqs = doc.add_paragraph()
    tech_reqs.add_run('Technical specifications include:').bold = True
    tech_reqs.add_run('\n• Microservices architecture\n• RESTful APIs\n• Database clustering\n• Load balancing')
    
    # Add security considerations
    doc.add_heading('Security Considerations', level=1)
    security = doc.add_paragraph()
    security.add_run('Security measures must include:').bold = True
    security.add_run('\n• Authentication and authorization\n• Data encryption at rest and in transit\n• Network security and firewalls\n• Regular security audits')
    
    # Add a table
    doc.add_heading('System Components', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Component'
    header_cells[1].text = 'Technology'
    header_cells[2].text = 'Purpose'
    
    # Add data rows
    components = [
        ('Frontend', 'React.js', 'User interface'),
        ('Backend', 'Node.js', 'Business logic'),
        ('Database', 'PostgreSQL', 'Data storage'),
        ('Cache', 'Redis', 'Performance optimization'),
        ('Load Balancer', 'Nginx', 'Traffic distribution')
    ]
    
    for component, tech, purpose in components:
        row_cells = table.add_row().cells
        row_cells[0].text = component
        row_cells[1].text = tech
        row_cells[2].text = purpose
    
    # Save the document
    doc.save('sample_architecture.docx')
    print("✅ Created sample_architecture.docx successfully!")
    
except ImportError:
    print("❌ python-docx not available. Install with: pip install python-docx")
except Exception as e:
    print(f"❌ Error creating Word document: {e}")
